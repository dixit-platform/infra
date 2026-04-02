from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/riteshdixit/Documents/dixitrit/infra")
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "generated-assets"
OUTPUT_PATH = DOCS_DIR / "devsecops-platform-progress-guide.docx"
FLOW_IMAGE = ASSETS_DIR / "pipeline-flow.png"

BLUE_DARK = "1F4E79"
BLUE = "2E75B6"
BLUE_LIGHT = "BDD7EE"
GREEN_LIGHT = "E2EFDA"
ORANGE_LIGHT = "FCE4D6"
GRAY = "595959"
GRAY_BORDER = "CCCCCC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>'))


def set_cell_text(cell, text, bold=False, color="000000", size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_para(p, text, size=11, bold=False, italic=False, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    return r


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_para(p, text, size=18, bold=True, color=BLUE_DARK, space_after=10)
    elif level == 2:
        set_para(p, text, size=14, bold=True, color=BLUE, space_after=8)
    else:
        set_para(p, text, size=12, bold=True, color=GRAY, space_after=6)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.2)
        set_para(p, f"• {item}", size=11, color="000000", space_after=3)


def add_simple_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        if widths:
            hdr[i].width = widths[i]
        set_cell_shading(hdr[i], BLUE)
        set_cell_text(hdr[i], h, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            text = value
            fill = None
            if isinstance(value, tuple):
                text, fill = value
            if fill:
                set_cell_shading(cells[i], fill)
            set_cell_text(cells[i], text, size=10)
    return table


def add_page_break(doc):
    doc.add_page_break()


def make_flow_image(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 540
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_big = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 34)
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 24)
        font_small = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 20)
    except Exception:
        font_big = ImageFont.load_default()
        font = ImageFont.load_default()
        font_small = ImageFont.load_default()

    draw.text((50, 30), "CI/CD Flow Overview", fill=f"#{BLUE_DARK}", font=font_big)

    boxes = [
        ((60, 140, 320, 260), "Feature Branch\nPush / PR", BLUE_LIGHT),
        ((420, 140, 680, 260), "CI\nfmt validate lint\nsecurity plan", GREEN_LIGHT),
        ((780, 140, 1040, 260), "Merge to\ndev / stage / main", ORANGE_LIGHT),
        ((1140, 140, 1400, 260), "CD\nplan apply validate\ndestroy", BLUE_LIGHT),
        ((1500, 140, 1740, 260), "Release Tags\npromote / destroy", GREEN_LIGHT),
    ]

    for (x1, y1, x2, y2), text, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=f"#{fill}", outline=f"#{GRAY_BORDER}", width=3)
        lines = text.split("\n")
        total_h = len(lines) * 28
        y = (y1 + y2) // 2 - total_h // 2
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            tw = bbox[2] - bbox[0]
            draw.text(((x1 + x2 - tw) / 2, y), line, fill="black", font=font)
            y += 30

    arrows = [
        (320, 200, 420, 200),
        (680, 200, 780, 200),
        (1040, 200, 1140, 200),
        (1400, 200, 1500, 200),
    ]
    for x1, y1, x2, y2 in arrows:
        draw.line((x1, y1, x2, y2), fill=f"#{BLUE_DARK}", width=6)
        draw.polygon([(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)], fill=f"#{BLUE_DARK}")

    footer = "Feature branches default to dev validation. Protected branches map to dev / stage / prod."
    draw.text((120, 390), footer, fill=f"#{GRAY}", font=font_small)
    img.save(path)


def build_doc():
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    make_flow_image(FLOW_IMAGE)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for style_name in ["Normal"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_para(p, "DevSecOps GitOps Platform", size=28, bold=True, color=BLUE_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    set_para(p, "Complete Setup & Learning Guide", size=16, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    set_para(
        p,
        "AWS · Terraform · GitHub Actions · DevSecOps · Platform Engineering",
        size=12,
        italic=True,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    headers = ["Steps Completed", "Resume From", "End Goal"]
    values = [
        ("Steps 1–11 completed", GREEN_LIGHT),
        ("Step 12: VPC Module + CI Alignment", ORANGE_LIGHT),
        ("Internal Developer Platform", BLUE_LIGHT),
    ]
    for i, h in enumerate(headers):
        set_cell_shading(table.cell(0, i), BLUE)
        set_cell_text(table.cell(0, i), h, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        text, fill = values[i]
        set_cell_shading(table.cell(1, i), fill)
        set_cell_text(table.cell(1, i), text, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    set_para(p, "Version 1.1 · Last updated: April 2, 2026", size=9, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)

    heading(doc, "Project Vision", 1)
    p = doc.add_paragraph()
    set_para(
        p,
        "Building a production-grade, multi-environment infrastructure platform on AWS using Terraform and GitHub Actions. "
        "The immediate goal is learning DevOps and DevSecOps at production quality. "
        "The long-term goal is an Internal Developer Platform where developers self-serve infrastructure through a portal.",
        size=11,
    )

    add_simple_table(
        doc,
        ["Phase", "What we are building"],
        [
            ["Phase 1 — Foundation", "AWS security, IAM, Terraform pipeline, self-hosted EC2 runner, multi-environment setup, ephemeral deployments, release tags"],
            ["Phase 2 — Platform Modules", "Reusable Terraform modules: VPC, EC2, ALB, ECS, EKS, Lambda with docs and secure defaults"],
            ["Phase 3 — IDP", "Developer portal, service catalog, self-service workflows, cost visibility, drift detection, visual workflows"],
        ],
    )

    heading(doc, "Full Project Roadmap", 1)
    add_simple_table(
        doc,
        ["Step", "Status", "What it covers"],
        [
            ["Step 1–11", ("Done", GREEN_LIGHT), "Security baseline, IAM, runner, bootstrap, CI/CD foundation, promote, destroy"],
            ["Step 12", ("In Progress", ORANGE_LIGHT), "Reusable Terraform VPC module wired into dev, stage, and prod"],
            ["Step 13", ("Next", BLUE_LIGHT), "Import the existing runner VPC into Terraform state"],
            ["Step 14", "Planned", "EC2 and ALB reusable modules"],
            ["Step 15", "Planned", "ECS reusable module"],
            ["Step 16", "Planned", "Migrate runner from EC2 to ECS Fargate"],
            ["Step 17", "Planned", "EKS reusable module"],
            ["Step 18", "Planned", "Migrate runner to EKS with ARC"],
            ["Step 19", "Planned", "Internal developer portal"],
        ],
    )

    heading(doc, "Current State Snapshot", 1)
    heading(doc, "Repository", 2)
    add_bullets(
        doc,
        [
            "Terraform environments exist for dev, stage, and prod.",
            "GitHub Actions workflows exist for CI, CD, promote, and destroy.",
            "A reusable VPC module exists in terraform/modules/vpc.",
            "Self-hosted GitHub Actions runner is running on EC2.",
        ],
    )
    heading(doc, "Current Implementation State", 2)
    add_bullets(
        doc,
        [
            "dev, stage, and prod now use a shared VPC module wrapper pattern.",
            "Environment-specific values are supplied through terraform.tfvars files.",
            "Outputs are exposed from each environment for later module integration.",
            "CI is being aligned to map branch intent to the correct Terraform environment.",
        ],
    )

    heading(doc, "Locked-In Decisions", 1)
    add_bullets(
        doc,
        [
            "AWS region is ap-south-1.",
            "Terraform baseline is >= 1.10.0 with S3 native locking.",
            "No DynamoDB table for state locking.",
            "No static AWS access keys.",
            "No NAT Gateway for this learning platform.",
            "dev stays cost-optimized.",
            "stage and prod use stronger observability defaults.",
            "Reusable Terraform modules are the main implementation pattern.",
        ],
    )

    heading(doc, "CI/CD Architecture", 1)
    p = doc.add_paragraph()
    set_para(
        p,
        "The pipeline follows a simple GitOps path: feature branches validate first, protected branches map to environment-specific validation, and CD remains limited to dev, stage, and main merges.",
        size=11,
    )
    doc.add_picture(str(FLOW_IMAGE), width=Inches(6.8))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "Workflow Notes", 1)
    add_simple_table(
        doc,
        ["Workflow", "Current Direction", "Best Practice Notes"],
        [
            ["CI", "Run on push to any branch and PR to dev/stage/main", "Feature branches default to dev validation; protected branches map to dev, stage, prod"],
            ["CD", "Run only on push to dev, stage, main", "Correct for promotion and deploy intent"],
            ["Promote", "Use release tags for stage/prod movement", "Manual confirmation for prod remains a good control"],
            ["Destroy", "Manual only with audit trail", "Strong safety posture; keep approvals and audit log"],
        ],
    )

    heading(doc, "Current Step", 1)
    heading(doc, "Step 12: First VPC Module", 2)
    add_bullets(
        doc,
        [
            "Replace old demo resources with a reusable VPC module in dev, stage, and prod.",
            "Expose environment outputs for future ECS, ALB, and EKS modules.",
            "Keep dev cheaper, stage/prod more production-like.",
        ],
    )
    add_simple_table(
        doc,
        ["Item", "Value / Decision"],
        [
            ["Dev VPC CIDR", "10.0.0.0/16"],
            ["Stage VPC CIDR", "10.1.0.0/16"],
            ["Prod VPC CIDR", "10.2.0.0/16"],
            ["AZ Count", "dev=1, stage=2, prod=2"],
            ["Flow Logs", "disabled in dev, enabled in stage and prod"],
            ["NAT Gateway", "not used"],
        ],
    )

    heading(doc, "Security & Compliance Fixes", 1)
    add_simple_table(
        doc,
        ["Checkov Finding", "Why it failed", "Fix / Decision"],
        [
            ["CKV2_AWS_12", "Default security group was unmanaged", "Default SG locked down explicitly"],
            ["CKV_AWS_130", "Public subnets auto-assigned public IPs", "map_public_ip_on_launch set to false"],
            ["CKV_AWS_158", "CloudWatch Log Group not encrypted with KMS", "Add customer-managed KMS key for flow logs"],
            ["CKV_AWS_338", "Retention less than 1 year", "Stage and prod flow-log retention should be >= 365 days"],
            ["CKV_AWS_290 / 355", "IAM policy was too broad", "Scope write permissions to the specific log group ARN"],
        ],
    )

    heading(doc, "Self-Hosted Runner Notes", 1)
    p = doc.add_paragraph()
    set_para(
        p,
        "Disk space pressure on the EC2 runner is expected over time. Common cleanup candidates are old runner tarballs, old versioned bin/externals directories, _diag logs, Docker artifacts, and tool caches.",
        size=11,
    )
    add_simple_table(
        doc,
        ["Runner Item", "Can Clean?", "Notes"],
        [
            ["actions-runner-linux-x64-*.tar.gz", ("Yes", GREEN_LIGHT), "Installer archive is safe to remove"],
            ["_diag old logs", ("Yes", GREEN_LIGHT), "Keep only recent logs if needed"],
            ["bin.older-version", ("Usually", ORANGE_LIGHT), "Remove only after confirming active version"],
            ["externals.older-version", ("Usually", ORANGE_LIGHT), "Remove only after confirming active version"],
            ["bin / externals current", ("No", BLUE_LIGHT), "These are active runner components"],
        ],
    )

    heading(doc, "Local Verification Practice", 1)
    add_bullets(
        doc,
        [
            "Run terraform fmt -check -recursive terraform before push.",
            "Run terraform init -backend=false and terraform validate for affected environments.",
            "Use CI for heavier checks like Checkov, Trivy, Gitleaks, and full backend-backed plan.",
            "If a shared module changes, validate dev, stage, and prod locally.",
        ],
    )

    heading(doc, "Working Notes", 1)
    add_simple_table(
        doc,
        ["Date", "Step", "Update"],
        [
            ["2026-04-01", "Step 12", "Replaced demo environment pattern with reusable VPC module wrappers for dev, stage, and prod"],
            ["2026-04-01", "Step 12", "Added environment-specific tfvars, outputs, and typed variables for all three environments"],
            ["2026-04-01", "Step 12", "Fixed VPC security findings for default security group and public subnet public-IP behavior"],
            ["2026-04-02", "CI", "CI workflow aligned conceptually to map branch intent to the correct Terraform environment"],
            ["2026-04-02", "Ops", "Reviewed self-hosted runner disk usage and identified safe cleanup candidates"],
            ["2026-04-02", "Security", "New Checkov findings identified around flow-log KMS, retention, and IAM scoping"],
        ],
    )

    heading(doc, "Next Step", 1)
    add_bullets(
        doc,
        [
            "Finalize ci.yml and .gitignore updates.",
            "Apply flow-log KMS, retention, and IAM scoping improvements in the VPC module.",
            "Push changes and observe CI results.",
            "Move to Step 13: import the existing runner VPC into Terraform state.",
        ],
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
